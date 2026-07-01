"""Generate sample PDF and DOCX files for testing AI document upload."""

from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent

CONTENT = {
    "title": "Photosynthesis — Grade 9 Science Review Sheet",
    "intro": (
        "Photosynthesis is the process used by plants, algae, and some bacteria to "
        "convert light energy into chemical energy stored in glucose. This sheet "
        "summarizes the main ideas students should know for a short quiz."
    ),
    "sections": [
        (
            "1. Definition and Purpose",
            [
                "Photosynthesis means 'making with light.'",
                "It produces food (glucose) for the organism and oxygen for other life.",
                "It occurs mainly in the chloroplasts of plant cells.",
            ],
        ),
        (
            "2. Chemical Equation",
            [
                "Overall equation: 6CO2 + 6H2O + light energy -> C6H12O6 + 6O2",
                "Reactants: carbon dioxide and water.",
                "Products: glucose and oxygen.",
            ],
        ),
        (
            "3. Two Main Stages",
            [
                "Light-dependent reactions occur in the thylakoid membranes.",
                "The Calvin cycle (light-independent reactions) occurs in the stroma.",
                "ATP and NADPH from the light stage power the Calvin cycle.",
            ],
        ),
        (
            "4. Factors That Affect the Rate",
            [
                "Light intensity — more light generally increases the rate up to a limit.",
                "Carbon dioxide concentration — more CO2 can increase the rate.",
                "Temperature — enzymes work best at moderate temperatures.",
                "Water availability — drought can slow or stop photosynthesis.",
            ],
        ),
        (
            "5. Key Vocabulary",
            [
                "Chlorophyll — green pigment that absorbs light.",
                "Stomata — small openings on leaves where gas exchange occurs.",
                "Autotroph — organism that makes its own food.",
                "Glucose — sugar produced during photosynthesis.",
            ],
        ),
        (
            "6. Sample Review Questions (for teachers)",
            [
                "What are the two main products of photosynthesis?",
                "Where in the chloroplast do light-dependent reactions occur?",
                "Name two factors that affect the rate of photosynthesis.",
                "What gas do plants take in for photosynthesis?",
                "True or False: Photosynthesis releases carbon dioxide as its main product.",
            ],
        ),
    ],
}


def build_plain_text():
    lines = [CONTENT["title"], "", CONTENT["intro"], ""]
    for heading, bullets in CONTENT["sections"]:
        lines.append(heading)
        for item in bullets:
            lines.append(f"  • {item}")
        lines.append("")
    return "\n".join(lines)


def create_docx(path: Path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(CONTENT["title"], level=1)
    doc.add_paragraph(CONTENT["intro"])
    for heading, bullets in CONTENT["sections"]:
        doc.add_heading(heading, level=2)
        for item in bullets:
            p = doc.add_paragraph(item, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(11)
    doc.save(path)


def create_pdf(path: Path):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, CONTENT["title"])
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, CONTENT["intro"])
    pdf.ln(4)

    for heading, bullets in CONTENT["sections"]:
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(0, 7, heading)
        pdf.set_font("Helvetica", size=11)
        for item in bullets:
            pdf.multi_cell(0, 6, f"- {item}")
        pdf.ln(2)

    pdf.output(path)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUT_DIR / "photosynthesis-quiz-source.docx"
    pdf_path = OUT_DIR / "photosynthesis-quiz-source.pdf"
    txt_path = OUT_DIR / "photosynthesis-quiz-source.txt"

    txt_path.write_text(build_plain_text(), encoding="utf-8")
    create_docx(docx_path)
    create_pdf(pdf_path)
    print(f"Created: {docx_path}")
    print(f"Created: {pdf_path}")
    print(f"Created: {txt_path}")


if __name__ == "__main__":
    main()
